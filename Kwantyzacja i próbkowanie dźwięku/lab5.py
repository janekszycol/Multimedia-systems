import soundfile as sf
import scipy.fftpack
from docx import Document
from docx.shared import Inches
from scipy.interpolate import interp1d
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO


def plotAudio(data, fs, axs, fsize=2 ** 8, TimeMargin=[0, 0.02]):
    # plt.subplot(2, 1, 1)
    axs[0].set_xlim(TimeMargin)
    axs[0].set_xlabel("Sekundy")
    # plt.plot(np.arange(0, data.shape[0]) / fs, data)
    axs[0].plot(np.arange(0, data.shape[0]) / fs, data)

    axs[1].set_xlabel("Czestotliwosc")
    axs[1].set_ylabel("Decybele")
    yf = scipy.fftpack.fft(data, fsize)
    # dodanie eps zeby nie bylo zera
    yf += np.finfo(np.float32).eps
    dbYf = 20 * np.log10(np.abs(yf[:fsize // 2]))
    x = np.arange(0, fs / 2, fs / fsize)
    axs[1].plot(x, dbYf)

    # plt.show()
    return np.max(dbYf), x[np.argmax(dbYf)]


def kwant(data, bit):
    ogType = data.dtype
    d = 2 ** (bit) - 1
    if np.issubdtype(data.dtype, np.floating):
        m = -1
        n = 1
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max
    DataF = data.astype(float)
    DataF = (DataF - m) / (n - m)

    DataF = DataF * d
    DataF = np.round(DataF)
    DataF = DataF / d

    DataF = DataF * (n - m)
    DataF = DataF + m

    # kwantyzacja na DataF
    return DataF.astype(ogType)


def decymacja(data, fs, n):
    newFs = fs // n
    decyData = data[::n].copy()
    return decyData, newFs


def interp(data, Fs, FsNew, method='cubic'):
    N = len(data)
    N1 = int(FsNew * len(data) / Fs)

    t = np.linspace(0, N / Fs, N)
    t1 = np.linspace(0, N / Fs, N1)

    metode_lin = interp1d(t, data)
    metode_nonlin = interp1d(t, data, kind=method)

    y_lin = metode_lin(t1).astype(data.dtype)
    y_nonlin = metode_nonlin(t1).astype(data.dtype)

    return y_lin, y_nonlin, t, t1


data, fs = sf.read('sing_high1.wav', dtype=np.int32)

s1 = np.round(np.linspace(0, 255, 255, dtype=np.uint8))
s2 = np.round(np.linspace(np.iinfo(np.int32).min, np.iinfo(np.int32).max, 1000, dtype=np.int32))
s3 = np.linspace(-1, 1, 10000)
s4 = kwant(s3, 2)
# plt.plot(s3,s4)
# plt.show()
decyData = decymacja(data, fs, 100)
# plt.plot(data)
# plt.plot(np.linspace(0,fs,len(decyData[0])),decyData[0])
# plt.show()
# fig, axs = plt.subplots(2, figsize=(10, 7))
# plotAudio(data,fs,axs)

# fig, axs = plt.subplots(2, figsize=(10, 7))
# dataInterpolatedLinear,dataInterpolatedNonLinear,t,t1=interp(data,fs,44100)
# axs[0].plot(t,data)
# axs[1].plot(t1,dataInterpolatedLinear)
# plt.show()


# sf.write('NowyDzwiek.wav', dataInterpolatedLinear, 44100)


# document = Document()
# document.add_heading('Wprowadzenie do pracy z dzwiekiem', 0)
# filesSin = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav', 'sin_combined.wav']
# filesSing = ['sing_high1.wav', 'sing_medium1.wav', 'sing_low1.wav']
# bits = [4, 8, 16, 24]
# steps = [2, 4, 6, 10, 24]
# Fsize = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
# document.add_heading('Kwantyfikacja', 1)
# for file in filesSin:
#     document.add_heading('Plik - {}'.format(file), 2)
#     data, fs = sf.read(file, dtype=np.int32)
#     TimeMargin = [0, 0.02]
#     if file == 'sin_8000Hz.wav':
#         TimeMargin = [0, 0.005]
#     else:
#         TimeMargin = [0, 0.02]
#
#     for i in range(len(bits)):
#         fig, axs = plt.subplots(2, figsize=(10, 7))
#         dataCopy = data.copy()
#         dataCopy = kwant(dataCopy, bits[i])
#         plotAudio(dataCopy, fs, axs, TimeMargin=TimeMargin)
#         fig.suptitle('Sygnał zapisany na ={}bitach'.format(str(bits[i])))  # Tytuł wykresu
#         fig.tight_layout(pad=1.5)  # poprawa czytelności
#         memfile = BytesIO()  # tworzenie bufora
#         fig.savefig(memfile)  # z zapis do bufora
#
#         document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
#
#         memfile.close()
#
# document.add_heading('Decymacja', 1)
# for file in filesSin:
#     document.add_heading('Plik - {}'.format(file), 2)
#     data, fs = sf.read(file, dtype=np.int32)
#     TimeMargin = [0, 0.02]
#     if file == 'sin_8000Hz.wav':
#         TimeMargin = [0, 0.005]
#     else:
#         TimeMargin = [0, 0.02]
#     for i in range(len(steps)):
#         fig, axs = plt.subplots(2, figsize=(10, 7))
#         dataCopy = data.copy()
#         dataCopy,newFs = decymacja(dataCopy, fs, steps[i])
#         plotAudio(dataCopy, newFs, axs, TimeMargin=TimeMargin)
#         fig.suptitle('Sygnał zapisany przy użyciu co {}-tej próbki'.format(str(steps[i])))  # Tytuł wykresu
#         fig.tight_layout(pad=1.5)  # poprawa czytelności
#         memfile = BytesIO()  # tworzenie bufora
#         fig.savefig(memfile)  # z zapis do bufora
#
#         document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
#
#         memfile.close()
#
# document.add_heading('Interpolacja', 1)
# for file in filesSin:
#     document.add_heading('Plik - {}'.format(file), 2)
#     data, fs = sf.read(file, dtype=np.int32)
#     TimeMargin = [0, 0.02]
#     if file == 'sin_8000Hz.wav':
#         TimeMargin = [0, 0.005]
#     else:
#         TimeMargin = [0, 0.02]
#     for i in range(len(Fsize)):
#
#         dataCopy = data.copy()
#         y_lin, y_nonlin, t, t1 = interp(dataCopy, fs, Fsize[i])
#
#         fig, axs = plt.subplots(2, figsize=(10, 7))
#         plotAudio(y_lin, Fsize[i], axs, TimeMargin=TimeMargin)
#         fig.suptitle('Interpolacja liniowa dla {}Hz'.format(str(Fsize[i])))  # Tytuł wykresu
#         fig.tight_layout(pad=1.5)  # poprawa czytelności
#         memfile = BytesIO()  # tworzenie bufora
#         fig.savefig(memfile)  # z zapis do bufora
#         document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
#         memfile.close()
#
#         fig, axs = plt.subplots(2, figsize=(10, 7))
#         plotAudio(y_nonlin, Fsize[i], axs, TimeMargin=TimeMargin)
#         fig.suptitle('Interpolacja liniowa dla {}Hz'.format(str(Fsize[i])))  # Tytuł wykresu
#         fig.tight_layout(pad=1.5)  # poprawa czytelności
#         memfile = BytesIO()  # tworzenie bufora
#         fig.savefig(memfile)  # z zapis do bufora
#         document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
#         memfile.close()
#
#
#
# document.save('Sprawozdanie lab 1.docx')  # zapis do pliku

filesSing = ['sing_high1.wav', 'sing_medium1.wav', 'sing_low1.wav']
bits = [4, 8]
steps = [4, 6, 10, 24]
Fsize = [4000, 8000, 11999, 16000, 16953]

for file in filesSing:
    for i in range(len(bits)):
        data, fs = sf.read(file, dtype=np.int32)
        newName=file.rsplit('.',1)
        dataCopy = data.copy()
        dataCopy = kwant(dataCopy, bits[i])
        sf.write(newName[0]+'bits{}'.format(str(bits[i]))+'.wav',dataCopy, fs)

for file in filesSing:
    for i in range(len(steps)):
        data, fs = sf.read(file, dtype=np.int32)
        newName=file.rsplit('.',1)
        dataCopy = data.copy()
        dataCopy,newFs = decymacja(dataCopy, fs, steps[i])
        sf.write(newName[0]+'steps{}'.format(str(steps[i]))+'.wav',dataCopy, newFs)

for file in filesSing:
    for i in range(len(Fsize)):
        data, fs = sf.read(file, dtype=np.int32)
        newName=file.rsplit('.',1)
        dataCopy = data.copy()
        y_lin, y_nonlin, t, t1 = interp(dataCopy, fs, Fsize[i])
        sf.write(newName[0]+'fs{}'.format(str(Fsize[i]))+'Linear.wav',y_lin, Fsize[i])
        sf.write(newName[0] + 'fs{}'.format(str(Fsize[i])) + 'NonLinear.wav', y_nonlin, Fsize[i])

