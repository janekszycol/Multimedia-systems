import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

pallet8 = np.array([
    [0.0, 0.0, 0.0, ],
    [0.0, 0.0, 1.0, ],
    [0.0, 1.0, 0.0, ],
    [0.0, 1.0, 1.0, ],
    [1.0, 0.0, 0.0, ],
    [1.0, 0.0, 1.0, ],
    [1.0, 1.0, 0.0, ],
    [1.0, 1.0, 1.0, ],
])

pallet16 = np.array([
    [0.0, 0.0, 0.0, ],
    [0.0, 1.0, 1.0, ],
    [0.0, 0.0, 1.0, ],
    [1.0, 0.0, 1.0, ],
    [0.0, 0.5, 0.0, ],
    [0.5, 0.5, 0.5, ],
    [0.0, 1.0, 0.0, ],
    [0.5, 0.0, 0.0, ],
    [0.0, 0.0, 0.5, ],
    [0.5, 0.5, 0.0, ],
    [0.5, 0.0, 0.5, ],
    [1.0, 0.0, 0.0, ],
    [0.75, 0.75, 0.75, ],
    [0.0, 0.5, 0.5, ],
    [1.0, 1.0, 1.0, ],
    [1.0, 1.0, 0.0, ]
])

M2 = np.array([
    [0, 8, 2, 10],
    [12, 4, 14, 6],
    [3, 11, 1, 9],
    [15, 7, 13, 5],
])


def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        return img
    else:
        img = img / 255.0
        return img


def colorFit(pixel, pallet):
    return pallet[np.argmin(np.linalg.norm(pallet - pixel, axis=1))]


def kwant_colorFit(img, Pallet):
    out_img = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            out_img[w, k] = colorFit(img[w, k], Pallet)
    return out_img


def kwant_colorFit_dithering(img, Pallet, Mpre, r):
    out_img = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            Cn = img[w, k] + r * Mpre[w % (2 * 2)][k % (2 * 2)]
            out_img[w, k] = colorFit(Cn, Pallet)
    return out_img


def floydSteinbergDithering(img, pallet):
    for x in range(img.shape[0]):
        for y in range(img.shape[1]):
            oldpixel = img[x, y]
            newpixel = colorFit(oldpixel, pallet)
            img[x, y] = newpixel
            quant_error = oldpixel - newpixel
            # img[x + 1, y] = img[x + 1, y] + quant_error * 7 / 16
            # img[x - 1, y + 1] = img[x - 1, y + 1] + quant_error * 3 / 16
            # img[x, y + 1] = img[x, y + 1] + quant_error * 5 / 16
            # img[x + 1, y + 1] = img[x + 1, y + 1] + quant_error * 1 / 16
            if x + 1 < img.shape[0]:
                img[x + 1, y] = img[x + 1, y] + quant_error * 7 / 16
            if x - 1 < img.shape[0] and y + 1 < img.shape[1]:
                img[x - 1, y + 1] = img[x - 1, y + 1] + quant_error * 3 / 16
            if y + 1 < img.shape[1]:
                img[x, y + 1] = img[x, y + 1] + quant_error * 5 / 16
            if x + 1 < img.shape[0] and y + 1 < img.shape[1]:
                img[x + 1, y + 1] = img[x + 1, y + 1] + quant_error * 1 / 16
    return img


def randomDithering(img):
    r = np.random.rand(img.shape[0], img.shape[1])
    if len(img.shape) > 2:
        for i in range(r.shape[0]):
            for j in range(r.shape[1]):
                if img[i][j][0] >= r[i][j]:
                    r[i][j] = 1
                else:
                    r[i][j] = 0
    else:
        for i in range(r.shape[0]):
            for j in range(r.shape[1]):
                if img[i][j] >= r[i][j]:
                    r[i][j] = 1
                else:
                    r[i][j] = 0
    return r


N = 2
paleta = np.linspace(0, 1, N).reshape(N, 1)
img = imgToFloat(plt.imread('GS_0003.png'))
# kwantyzacja 1 czarno biale
#
# fig, axs = plt.subplots(1, 4)
# axs[0].imshow(img, cmap=plt.cm.gray)
# for i in range(1, 4):
#     paleta = np.linspace(0, 1, 2 ** i).reshape(2 ** i, 1)
#     kwantImg = kwant_colorFit(img, paleta)
#     axs[i].imshow(kwantImg, cmap=plt.cm.gray)
#
# plt.show()
#
# # #kwantyzacja 2 kolorowe
# palety = [pallet8, pallet16]
#
# img = imgToFloat(plt.imread('SMALL_0006.jpg'))
# fig, axs = plt.subplots(1, 3)
# axs[0].imshow(img)
# for i in range(1, 3):
#     kwantImg = kwant_colorFit(img, palety[i - 1])
#     axs[i].imshow(kwantImg)
#
# plt.show()
#
# # # dithering losowy
# r = randomDithering(img.copy())
# plt.imshow(r, cmap=plt.cm.gray)
# plt.show()
#
# # dithering zorganizowany
# Mpre = (M2 + 1) / (2 * 2) ** 2 - 0.5
# organizedDithering = kwant_colorFit_dithering(img.copy(), pallet16, Mpre, 1)
# plt.imshow(organizedDithering)
# plt.show()
#
# # dithering Floyda Steinberga
# r = floydSteinbergDithering(img.copy(), pallet16)
# plt.imshow(r, cmap=plt.cm.gray)
# plt.show()


files = ['GS_0001.tif', 'GS_0002.png', 'GS_0003.png', 'SMALL_0001.tif', 'SMALL_0005.jpg', 'SMALL_0006.jpg',
         'SMALL_0007.jpg']

filesGS = ['GS_0001.tif', 'GS_0002.png', 'GS_0003.png']
filesC = ['SMALL_0001.tif', 'SMALL_0005.jpg', 'SMALL_0006.jpg', 'SMALL_0007.jpg']

for file in files:
    img = imgToFloat(plt.imread(file))
    print(img.shape)
# fragments = [[[0, 1120, 200, 1320], [200, 400, 400, 600]], [[0, 0, 200, 200], [400, 400, 600, 600]]]

document = Document()
document.add_heading('Dithering i kwantyzacja obrazu', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka
for file in files:
    if file in filesGS:
        for i in range(1, 4):
            document.add_heading('Dithering i kwantyzacja obrazu - {}bits'.format(i), 2)
            paleta = np.linspace(0, 1, 2 ** i).reshape(2 ** i, 1)
            if i == 1:
                img = imgToFloat(plt.imread(file))
                if len(img.shape)>2:
                    R = img[:, :, 0]
                    G = img[:, :, 1]
                    B = img[:, :, 2]
                    img = 0.2126 * R + 0.7152 * G + 0.0722 * B

                fig = plt.figure(figsize=(8, 11))

                plt.subplot(3, 2, 1, title="Oryginał")
                plt.imshow(img, cmap=plt.cm.gray)

                kwantImg = kwant_colorFit(img.copy(), paleta)
                plt.subplot(3, 2, 2, title="Kwantyzacja")
                plt.imshow(kwantImg, cmap=plt.cm.gray)

                randDithering = randomDithering(img.copy())
                plt.subplot(3, 2, 3, title="Dithering losowy")
                plt.imshow(randDithering, cmap=plt.cm.gray)

                Mpre = (M2 + 1) / (2 * 2) ** 2 - 0.5
                organizedDithering = kwant_colorFit_dithering(img.copy(), paleta, Mpre, 1)
                plt.subplot(3, 2, 4, title="Dithering zorganizowany")
                plt.imshow(organizedDithering, cmap=plt.cm.gray)

                floydSteinbergDith = floydSteinbergDithering(img.copy(), paleta)
                plt.subplot(3, 2, 5, title="Dithering Floyda-Steinberga")
                plt.imshow(floydSteinbergDith, cmap=plt.cm.gray)

                fig.tight_layout(pad=1.5)  # poprawa czytelności
                memfile = BytesIO()  # tworzenie bufora
                fig.savefig(memfile)  # z zapis do bufora

                document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

                memfile.close()
            else:
                img = imgToFloat(plt.imread(file))
                fig = plt.figure(figsize=(8, 11))

                plt.subplot(2, 2, 1, title="Oryginał")
                plt.imshow(img, cmap=plt.cm.gray)

                kwantImg = kwant_colorFit(img.copy(), paleta)
                plt.subplot(2, 2, 2, title="Kwantyzacja")
                plt.imshow(kwantImg, cmap=plt.cm.gray)

                Mpre = (M2 + 1) / (2 * 2) ** 2 - 0.5
                organizedDithering = kwant_colorFit_dithering(img.copy(), paleta, Mpre, 1)
                plt.subplot(2, 2, 3, title="Dithering zorganizowany")
                plt.imshow(organizedDithering, cmap=plt.cm.gray)

                floydSteinbergDith = floydSteinbergDithering(img.copy(), paleta)
                plt.subplot(2, 2, 4, title="Dithering Floyda-Steinberga")
                plt.imshow(floydSteinbergDith, cmap=plt.cm.gray)

                fig.tight_layout(pad=1.5)  # poprawa czytelności
                memfile = BytesIO()  # tworzenie bufora
                fig.savefig(memfile)  # z zapis do bufora

                document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

                memfile.close()
    else:
        pallets = [pallet8, pallet16]
        for pallet in pallets:
            document.add_heading('Dithering i kwantyzacja obrazu - {}'.format(pallet), 2)
            img = imgToFloat(plt.imread(file))
            fig = plt.figure(figsize=(8, 11))

            plt.subplot(2, 2, 1, title="Oryginał")
            plt.imshow(img)

            kwantImg = kwant_colorFit(img.copy(), pallet)
            plt.subplot(2, 2, 2, title="Kwantyzacja")
            plt.imshow(kwantImg)

            Mpre = (M2 + 1) / (2 * 2) ** 2 - 0.5
            organizedDithering = kwant_colorFit_dithering(img.copy(), pallet, Mpre, 1)
            plt.subplot(2, 2, 3, title="Dithering zorganizowany")
            plt.imshow(organizedDithering)

            floydSteinbergDith = floydSteinbergDithering(img.copy(), pallet)
            plt.subplot(2, 2, 4, title="Dithering Floyda-Steinberga")
            plt.imshow(floydSteinbergDith)

            fig.tight_layout(pad=1.5)  # poprawa czytelności
            memfile = BytesIO()  # tworzenie bufora
            fig.savefig(memfile)  # z zapis do bufora

            document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

            memfile.close()

    # fig, axs = plt.subplots(2, 2, figsize=(10, 7))  # tworzenie plota
    # img = imgToFloat(plt.imread(file))
    # ############################################################
    # # Tu wykonujesz jakieś funkcje i rysujesz wykresy
    # fig, axs = plt.subplots(1, 4, figsize=(10, 7))
    # axs[0].imshow(img, cmap=plt.cm.gray)
    # for i in range(1, 4):
    #     paleta = np.linspace(0, 1, 2 ** i).reshape(2 ** i, 1)
    #     kwantImg = kwant_colorFit(img, paleta)
    #     axs[i].set_title("")
    #     axs[i].imshow(kwantImg, cmap=plt.cm.gray)
    # ############################################################
    #
    #  # fig.suptitle('Time margin {}'.format())  # Tytuł wykresu
    # fig.tight_layout(pad=1.5)  # poprawa czytelności
    # memfile = BytesIO()  # tworzenie bufora
    # fig.savefig(memfile)  # z zapis do bufora
    #
    # document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
    #
    # memfile.close()
    ############################################################
    # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
    ############################################################

document.save('Sprawozdanie dithering i kwantyzacja obrazu.docx')  # zapis do pliku
